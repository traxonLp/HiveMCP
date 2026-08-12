"""Render a DocSpec into a .docx file."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .base import MEDIA_TYPES, RenderedFile, RenderError
from .theme import (
    DEFAULT_MONO_FONT,
    ImageResolver,
    RenderWarnings,
    check_font,
    decode_image,
    estimate_pages,
    image_size,
    normalize_hex,
    safe_filename,
)

PAGE_SIZES_CM = {"A4": (21.0, 29.7), "Letter": (21.59, 27.94)}

ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

PARAGRAPH_STYLES = {
    "normal": ("Normal",),
    "quote": ("Quote", "Intense Quote", "Normal"),
    "caption": ("Caption", "Normal"),
    "intense": ("Intense Quote", "Quote", "Normal"),
}


class DocxRenderer:
    def __init__(
        self,
        options: Any,
        image_resolver: ImageResolver | None = None,
        template_path: Path | None = None,
    ) -> None:
        self.options = options
        self.image_resolver = image_resolver
        self.template_path = template_path
        self.warnings = RenderWarnings()
        self.font = check_font(getattr(options, "font_family", None), self.warnings)
        self._character_count = 0

    # ---------------------------------------------------------------- public

    def render(self, spec: Any) -> RenderedFile:
        document = self._open_document()
        self._apply_page_setup(document)
        self._apply_base_style(document)
        self._apply_metadata(document, spec)

        self._add_front_matter(document, spec)

        for index, block in enumerate(spec.blocks):
            handler = getattr(self, f"_block_{block.type}", None)
            if handler is None:
                raise RenderError(f"block {index + 1}: unknown type {block.type!r}")
            try:
                handler(document, block)
            except Exception as exc:  # noqa: BLE001 - a bad block must name itself
                raise RenderError(
                    f"block {index + 1} ({block.type}) could not be rendered: {exc}"
                ) from exc

        if getattr(spec, "placeholders", None):
            self._fill_placeholders(document, spec.placeholders)

        buffer = BytesIO()
        document.save(buffer)
        return RenderedFile(
            data=buffer.getvalue(),
            filename=safe_filename(getattr(self.options, "filename", None), spec.title, "docx"),
            media_type=MEDIA_TYPES["docx"],
            warnings=list(self.warnings),
            page_estimate=estimate_pages(
                self._character_count, getattr(self.options, "density", "normal")
            ),
        )

    # --------------------------------------------------------------- private

    def _open_document(self) -> Any:
        if self.template_path is None:
            return Document()
        try:
            return Document(str(self.template_path))
        except Exception as exc:  # noqa: BLE001
            raise RenderError(
                f"template {self.template_path.name!r} could not be opened as a Word "
                f"file: {exc}"
            ) from exc

    def _style(self, document: Any, candidates: tuple[str, ...]) -> str | None:
        """Return the first style the document actually defines.

        Custom templates routinely lack 'List Bullet 3' or 'Quote'. Falling back
        quietly beats raising, but the caller should know the document will not look
        exactly as asked.
        """
        available = {style.name for style in document.styles}
        for candidate in candidates:
            if candidate in available:
                return candidate
        if candidates:
            self.warnings.add(
                f"Template defines none of the styles {list(candidates)}; used default "
                "paragraph formatting instead."
            )
        return None

    def _apply_page_setup(self, document: Any) -> None:
        width_cm, height_cm = PAGE_SIZES_CM[getattr(self.options, "page_size", "A4")]
        landscape = getattr(self.options, "orientation", "portrait") == "landscape"
        if landscape:
            width_cm, height_cm = height_cm, width_cm
        for section in document.sections:
            section.page_width = Cm(width_cm)
            section.page_height = Cm(height_cm)
            section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT

    def _apply_base_style(self, document: Any) -> None:
        """Set font on the Normal style so every derived style inherits it."""
        if not self.font and not getattr(self.options, "font_size_base", None):
            return
        normal = document.styles["Normal"]
        if self.font:
            normal.font.name = self.font
            # Word looks at the east-asian typeface for a surprising amount of Latin
            # text; without this the name silently reverts for some content.
            rpr = normal.element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:eastAsia"), self.font)
        if getattr(self.options, "font_size_base", None):
            normal.font.size = Pt(self.options.font_size_base)

    def _apply_metadata(self, document: Any, spec: Any) -> None:
        properties = document.core_properties
        properties.title = spec.title
        if getattr(spec, "author", None):
            properties.author = spec.author
        properties.language = getattr(self.options, "language", "de")

    def _add_front_matter(self, document: Any, spec: Any) -> None:
        title_style = self._style(document, ("Title", "Heading 1"))
        paragraph = document.add_paragraph(spec.title, style=title_style)
        self._count(spec.title)
        if getattr(spec, "subtitle", None):
            subtitle_style = self._style(document, ("Subtitle", "Normal"))
            document.add_paragraph(spec.subtitle, style=subtitle_style)
            self._count(spec.subtitle)
        _ = paragraph

    def _count(self, text: str | None) -> None:
        if text:
            self._character_count += len(text)

    # -- block handlers ----------------------------------------------------

    def _block_heading(self, document: Any, block: Any) -> None:
        document.add_heading(block.text, level=block.level)
        self._count(block.text)

    def _block_paragraph(self, document: Any, block: Any) -> None:
        style = self._style(document, PARAGRAPH_STYLES[block.style])
        paragraph = document.add_paragraph(block.text, style=style)
        paragraph.alignment = ALIGNMENTS[block.alignment]
        self._count(block.text)

    def _block_bullet_list(self, document: Any, block: Any) -> None:
        self._write_list(document, block.items, "List Bullet")

    def _block_numbered_list(self, document: Any, block: Any) -> None:
        self._write_list(document, block.items, "List Number")

    def _write_list(self, document: Any, items: list[Any], base: str, level: int = 0) -> None:
        # Word's built-in list styles only go three deep; deeper nesting reuses the
        # last style and relies on the indent below to stay readable.
        suffix = "" if level == 0 else f" {min(level + 1, 3)}"
        style = self._style(document, (f"{base}{suffix}", base, "List Paragraph", "Normal"))
        for item in items:
            paragraph = document.add_paragraph(style=style)
            if level >= 3:
                paragraph.paragraph_format.left_indent = Cm(0.75 * (level + 1))
            run = paragraph.add_run(item.text)
            run.bold = item.bold or None
            run.italic = item.italic or None
            self._count(item.text)
            if item.children:
                self._write_list(document, item.children, base, level + 1)

    def _block_table(self, document: Any, block: Any) -> None:
        data = block.data
        style = self._style(document, ("Table Grid", "Normal Table"))
        table = document.add_table(rows=1, cols=len(data.headers))
        if style:
            table.style = style

        for index, header in enumerate(data.headers):
            cell = table.rows[0].cells[index]
            cell.text = header
            self._count(header)
            if data.header_bold:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        for row in data.rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = str(value)
                self._count(str(value))

        if data.column_widths_cm:
            # Word ignores column width unless it is set on every cell of the column.
            for index, width_cm in enumerate(data.column_widths_cm):
                for row in table.rows:
                    row.cells[index].width = Cm(width_cm)

        if block.caption:
            caption_style = self._style(document, ("Caption", "Normal"))
            document.add_paragraph(block.caption, style=caption_style)
            self._count(block.caption)

    def _block_image(self, document: Any, block: Any) -> None:
        stream = decode_image(block.image, self.image_resolver)
        kwargs = image_size(block.image)
        if not kwargs:
            section = document.sections[0]
            usable = section.page_width - section.left_margin - section.right_margin
            kwargs["width"] = usable
        document.add_picture(stream, **kwargs)
        if block.caption:
            caption_style = self._style(document, ("Caption", "Normal"))
            document.add_paragraph(block.caption, style=caption_style)
            self._count(block.caption)

    def _block_page_break(self, document: Any, block: Any) -> None:
        document.add_page_break()

    def _block_code(self, document: Any, block: Any) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.5)
        run = paragraph.add_run(block.text)
        run.font.name = DEFAULT_MONO_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        self._shade(paragraph, "F2F2F2")
        self._count(block.text)

    def _block_toc(self, document: Any, block: Any) -> None:
        """Insert a real TOC field.

        A field, not a rendered list: Word computes page numbers itself and keeps the
        entries in sync as the document is edited. The trade-off is that the TOC shows
        placeholder text until the field is updated, so updateFields is set below to
        make Word offer that on open.
        """
        paragraph = document.add_paragraph()
        run = paragraph.add_run()

        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = f'TOC \\o "1-{block.depth}" \\h \\z \\u'
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = "Inhaltsverzeichnis - im Dokument mit F9 aktualisieren."
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")

        for element in (begin, instruction, separate, placeholder, end):
            run._r.append(element)  # noqa: SLF001

        self._request_field_update(document)

    @staticmethod
    def _request_field_update(document: Any) -> None:
        settings = document.settings.element
        if settings.find(qn("w:updateFields")) is not None:
            return
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)

    @staticmethod
    def _shade(paragraph: Any, hex_color: str) -> None:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), normalize_hex(hex_color) or "F2F2F2")
        paragraph._p.get_or_add_pPr().append(shading)  # noqa: SLF001

    @staticmethod
    def _fill_placeholders(document: Any, values: dict[str, str]) -> None:
        """Replace ``{{name}}`` tokens in body, tables, headers and footers.

        Replacement happens per run, so a token split across runs by Word's spell
        checker will not match. That is a known limit of run-level replacement; the
        docxtpl path used for uploaded templates handles it properly.
        """

        def replace_in_paragraph(paragraph: Any) -> None:
            for run in paragraph.runs:
                for key, value in values.items():
                    token = "{{" + key + "}}"
                    if token in run.text:
                        run.text = run.text.replace(token, value)

        def walk(container: Any) -> None:
            for paragraph in container.paragraphs:
                replace_in_paragraph(paragraph)
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        walk(cell)

        walk(document)
        for section in document.sections:
            walk(section.header)
            walk(section.footer)


def render_document(
    spec: Any,
    options: Any,
    image_resolver: ImageResolver | None = None,
    template_path: Path | None = None,
) -> RenderedFile:
    return DocxRenderer(options, image_resolver, template_path).render(spec)
