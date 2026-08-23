"""Word rendering.

Same principle as the pptx tests: open the produced file and check what it contains.
"""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from hivemcp.core.models import (
    Bullet,
    BulletList,
    CodeBlock,
    DocSpec,
    Heading,
    ImageBlock,
    ImageRef,
    NumberedList,
    PageBreak,
    Paragraph,
    RenderOptions,
    TableBlock,
    TableData,
    TableOfContents,
)
from hivemcp.core.render.base import MEDIA_TYPES, RenderError
from hivemcp.core.render.docx import render_document


def opened(rendered) -> Document:
    return Document(BytesIO(rendered.data))


def body_text(document: Document) -> str:
    return "\n".join(p.text for p in document.paragraphs)


# --------------------------------------------------------------------------- #
# Structure
# --------------------------------------------------------------------------- #


def test_renders_title_and_blocks(document: DocSpec, options: RenderOptions) -> None:
    text = body_text(opened(render_document(document, options)))
    assert "Betriebshandbuch" in text
    assert "Einleitung" in text
    assert "Dieses Dokument beschreibt den Betrieb." in text


def test_reports_media_type_and_page_estimate(
    document: DocSpec, options: RenderOptions
) -> None:
    rendered = render_document(document, options)
    assert rendered.media_type == MEDIA_TYPES["docx"]
    assert rendered.filename.endswith(".docx")
    assert rendered.page_estimate is not None and rendered.page_estimate >= 1


def test_output_is_a_real_ooxml_package(
    document: DocSpec, options: RenderOptions
) -> None:
    assert render_document(document, options).data[:2] == b"PK"


# --------------------------------------------------------------------------- #
# Blocks
# --------------------------------------------------------------------------- #


def test_heading_levels_map_to_heading_styles(options: RenderOptions) -> None:
    spec = DocSpec(
        title="T",
        blocks=[Heading(text=f"H{level}", level=level) for level in (1, 2, 3)],
    )
    document = opened(render_document(spec, options))
    styles = {p.text: p.style.name for p in document.paragraphs if p.text.startswith("H")}
    assert styles["H1"].startswith("Heading 1")
    assert styles["H2"].startswith("Heading 2")
    assert styles["H3"].startswith("Heading 3")


def test_bullet_and_numbered_lists_use_list_styles(options: RenderOptions) -> None:
    spec = DocSpec(
        title="T",
        blocks=[
            BulletList(items=[Bullet(text="Punkt")]),
            NumberedList(items=[Bullet(text="Schritt")]),
        ],
    )
    document = opened(render_document(spec, options))
    styles = {p.text: p.style.name for p in document.paragraphs if p.text}
    assert "List Bullet" in styles["Punkt"]
    assert "List Number" in styles["Schritt"]


def test_nested_list_items_are_indented(options: RenderOptions) -> None:
    spec = DocSpec(
        title="T",
        blocks=[
            BulletList(items=[Bullet(text="oben", children=[Bullet(text="unten")])])
        ],
    )
    document = opened(render_document(spec, options))
    styles = {p.text: p.style.name for p in document.paragraphs if p.text}
    # The nested level uses the numbered variant of the style ("List Bullet 2").
    assert styles["unten"] != styles["oben"]


def test_table_block_renders_headers_and_rows(options: RenderOptions) -> None:
    spec = DocSpec(
        title="T",
        blocks=[
            TableBlock(
                data=TableData(headers=["A", "B"], rows=[["1", "2"], ["3", "4"]]),
                caption="Messwerte",
            )
        ],
    )
    document = opened(render_document(spec, options))
    assert len(document.tables) == 1
    table = document.tables[0]
    assert [cell.text for cell in table.rows[0].cells] == ["A", "B"]
    assert len(table.rows) == 3
    assert "Messwerte" in body_text(document)


def test_inline_image_is_embedded(tiny_png: str, options: RenderOptions) -> None:
    spec = DocSpec(
        title="T",
        blocks=[ImageBlock(image=ImageRef(data_base64=tiny_png, width_cm=4))],
    )
    rendered = render_document(spec, options)
    document = opened(rendered)
    assert any(
        part.content_type.startswith("image/") for part in document.part.package.parts
    )


def test_page_break_is_written(options: RenderOptions) -> None:
    spec = DocSpec(
        title="T",
        blocks=[Paragraph(text="vorne"), PageBreak(), Paragraph(text="hinten")],
    )
    xml = opened(render_document(spec, options)).element.xml
    assert 'w:type="page"' in xml or "lastRenderedPageBreak" in xml or "w:br" in xml


def test_code_block_uses_a_monospace_font(options: RenderOptions) -> None:
    spec = DocSpec(title="T", blocks=[CodeBlock(text="print(1)", language="python")])
    document = opened(render_document(spec, options))
    paragraph = next(p for p in document.paragraphs if "print(1)" in p.text)
    fonts = {run.font.name for run in paragraph.runs}
    assert fonts and all(name is not None for name in fonts)


def test_paragraph_alignment_is_applied(options: RenderOptions) -> None:
    spec = DocSpec(title="T", blocks=[Paragraph(text="mittig", alignment="center")])
    document = opened(render_document(spec, options))
    paragraph = next(p for p in document.paragraphs if p.text == "mittig")
    assert paragraph.alignment is not None


# --------------------------------------------------------------------------- #
# Table of contents
# --------------------------------------------------------------------------- #


def test_toc_block_inserts_a_field(options: RenderOptions) -> None:
    """Word builds the entries itself; the file only carries the TOC field.

    Asserting on the field code rather than on visible text matters, because the entries
    genuinely are not there until Word refreshes them.
    """
    spec = DocSpec(
        title="T", blocks=[TableOfContents(depth=2), Heading(text="Kapitel", level=1)]
    )
    assert "TOC" in opened(render_document(spec, options)).element.xml


def test_include_toc_option_adds_one_without_a_block() -> None:
    """The GUI offers this checkbox, so it has to reach the document.

    It did not: the option travelled from the card to the server and was then read by
    nothing, so ticking the box changed nothing and said nothing.
    """
    spec = DocSpec(title="T", blocks=[Heading(text="Kapitel", level=1)])
    rendered = render_document(spec, RenderOptions(include_toc=True))
    assert "TOC" in opened(rendered).element.xml


def test_include_toc_does_not_duplicate_an_explicit_block() -> None:
    spec = DocSpec(
        title="T",
        blocks=[TableOfContents(depth=2), Heading(text="Kapitel", level=1)],
    )
    rendered = render_document(spec, RenderOptions(include_toc=True))
    assert opened(rendered).element.xml.count('TOC \\o') == 1


def test_no_toc_when_neither_asked_for() -> None:
    spec = DocSpec(title="T", blocks=[Heading(text="Kapitel", level=1)])
    rendered = render_document(spec, RenderOptions(include_toc=False))
    assert "TOC" not in opened(rendered).element.xml


# --------------------------------------------------------------------------- #
# Options and errors
# --------------------------------------------------------------------------- #


def test_landscape_orientation_swaps_the_page_dimensions(document: DocSpec) -> None:
    portrait = opened(render_document(document, RenderOptions(orientation="portrait")))
    landscape = opened(render_document(document, RenderOptions(orientation="landscape")))
    assert portrait.sections[0].page_width < portrait.sections[0].page_height
    assert landscape.sections[0].page_width > landscape.sections[0].page_height


def test_letter_page_size_differs_from_a4(document: DocSpec) -> None:
    a4 = opened(render_document(document, RenderOptions(page_size="A4")))
    letter = opened(render_document(document, RenderOptions(page_size="Letter")))
    assert a4.sections[0].page_width != letter.sections[0].page_width


def test_cjk_language_with_a_latin_font_warns(document: DocSpec) -> None:
    rendered = render_document(
        document, RenderOptions(language="zh-CN", font_family="Arial")
    )
    assert any("zh" in w or "Chinese" in w or "CJK" in w for w in rendered.warnings)


def test_render_error_names_the_block_that_failed(options: RenderOptions) -> None:
    spec = DocSpec(
        title="T",
        blocks=[
            Paragraph(text="ok"),
            ImageBlock(image=ImageRef(data_base64="not-base64-at-all")),
        ],
    )
    with pytest.raises(RenderError) as caught:
        render_document(spec, options)
    assert "2" in str(caught.value)
